import fitz  # PyMuPDF
from docx import Document
import re
from pathlib import Path


class ResumeParser:
    def parse(self, file_path: str) -> dict:
        path = Path(file_path)
        if path.suffix.lower() == ".pdf":
            text = self._parse_pdf(file_path)
        elif path.suffix.lower() in [".docx", ".doc"]:
            text = self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported format: {path.suffix}")

        return {
            "raw_text": text,
            "file_name": path.name,
            "word_count": len(text.split()),
            "char_count": len(text),
        }

    def _parse_pdf(self, file_path: str) -> str:
        doc = fitz.open(file_path)
        all_text = []

        for page in doc:
            # Method 1: Normal text extraction
            text = page.get_text("text")

            # Method 2: যদি text খুব কম আসে — layout-aware try করো
            if len(text.strip()) < 50:
                text = page.get_text("blocks")
                if isinstance(text, list):
                    text = "\n".join(b[4] for b in text if isinstance(b[4], str))

            # Method 3: Multi-column — dict mode দিয়ে reading order ঠিক করো
            if len(text.strip()) < 50:
                blocks = page.get_text("dict")["blocks"]
                lines = []
                for block in blocks:
                    if block.get("type") == 0:  # text block
                        for line in block.get("lines", []):
                            line_text = " ".join(
                                span["text"] for span in line.get("spans", [])
                            )
                            if line_text.strip():
                                lines.append(line_text)
                text = "\n".join(lines)

            all_text.append(text)

        doc.close()
        return "\n".join(all_text).strip()

    def _parse_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        parts = []

        # Normal paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables (অনেক resume table-based layout ব্যবহার করে)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)


class ResumeTextCleaner:
    def clean(self, text: str) -> str:
        if not text:
            return ""

        # Unicode normalize করো
        text = text.encode("utf-8", errors="ignore").decode("utf-8")

        # Extra whitespace
        text = re.sub(r"[ \t]+", " ", text)

        # ৩+ newline কমাও
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Bullet symbols normalize করো
        text = re.sub(r"[•·▪▸►◦‣⁃]", "-", text)

        # Special chars clean (email/phone রাখো)
        text = re.sub(r"[^\w\s\.\,\!\?\@\#\-\+\/\(\)\:\|]", " ", text)

        # Multiple spaces
        text = re.sub(r" {2,}", " ", text)

        return text.strip()
